"""
Export service for multiple formats.

Supports:
- PDF (formatted reports)
- Word (.docx)
- Markdown
- JSON
- CSV
"""

import io
import csv
import json
from datetime import datetime
from typing import List, Dict, Any, Optional
from enum import Enum

from sqlalchemy.orm import Session

from models import Meeting, Conversation, ActionItem, MeetingSummary


class ExportFormat(str, Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    JSON = "json"
    CSV = "csv"
    HTML = "html"


class ExportService:
    """
    Service for exporting meetings and data in various formats.
    """

    def __init__(self, db: Session):
        self.db = db

    def export_meeting(
        self,
        meeting: Meeting,
        format: ExportFormat,
        include_conversations: bool = True,
        include_action_items: bool = True,
        include_summary: bool = True,
    ) -> tuple:
        """
        Export a single meeting.

        Returns:
            Tuple of (content, filename, content_type)
        """
        # Gather meeting data
        data = self._gather_meeting_data(
            meeting,
            include_conversations,
            include_action_items,
            include_summary,
        )

        if format == ExportFormat.PDF:
            return self._export_pdf(data)
        elif format == ExportFormat.DOCX:
            return self._export_docx(data)
        elif format == ExportFormat.MARKDOWN:
            return self._export_markdown(data)
        elif format == ExportFormat.JSON:
            return self._export_json(data)
        elif format == ExportFormat.CSV:
            return self._export_csv(data)
        elif format == ExportFormat.HTML:
            return self._export_html(data)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def export_multiple_meetings(
        self,
        meetings: List[Meeting],
        format: ExportFormat,
    ) -> tuple:
        """
        Export multiple meetings.

        Returns:
            Tuple of (content, filename, content_type)
        """
        all_data = [self._gather_meeting_data(m) for m in meetings]

        if format == ExportFormat.JSON:
            content = json.dumps(all_data, indent=2, default=str)
            return (
                content.encode("utf-8"),
                f"meetings_export_{datetime.utcnow().strftime('%Y%m%d')}.json",
                "application/json",
            )
        elif format == ExportFormat.CSV:
            return self._export_meetings_csv(all_data)
        elif format == ExportFormat.MARKDOWN:
            return self._export_meetings_markdown(all_data)
        else:
            raise ValueError(f"Bulk export not supported for format: {format}")

    def _gather_meeting_data(
        self,
        meeting: Meeting,
        include_conversations: bool = True,
        include_action_items: bool = True,
        include_summary: bool = True,
    ) -> Dict[str, Any]:
        """Gather all meeting data for export."""
        data = {
            "id": meeting.id,
            "title": meeting.title or "Untitled Meeting",
            "meeting_type": meeting.meeting_type,
            "meeting_app": meeting.meeting_app,
            "started_at": meeting.started_at.isoformat() if meeting.started_at else None,
            "ended_at": meeting.ended_at.isoformat() if meeting.ended_at else None,
            "duration_minutes": meeting.duration_seconds // 60 if meeting.duration_seconds else 0,
            "status": meeting.status,
            "notes": meeting.notes,
        }

        if include_conversations:
            data["conversations"] = [
                {
                    "speaker": c.speaker,
                    "heard_text": c.heard_text,
                    "response_text": c.response_text,
                    "timestamp": c.timestamp.isoformat() if c.timestamp else None,
                }
                for c in meeting.conversations
            ]

        if include_action_items:
            data["action_items"] = [
                {
                    "description": a.description,
                    "assignee": a.assignee,
                    "due_date": a.due_date.isoformat() if a.due_date else None,
                    "status": a.status,
                    "priority": a.priority,
                }
                for a in meeting.action_items
            ]

        if include_summary and meeting.summary:
            data["summary"] = {
                "text": meeting.summary.summary_text,
                "key_points": meeting.summary.key_points,
                "decisions": meeting.summary.decisions_made,
                "sentiment": meeting.summary.sentiment,
            }

        return data

    def _export_markdown(self, data: Dict[str, Any]) -> tuple:
        """Export as Markdown."""
        lines = []

        # Header
        lines.append(f"# {data['title']}")
        lines.append("")
        lines.append(f"**Date:** {data['started_at']}")
        lines.append(f"**Duration:** {data['duration_minutes']} minutes")
        lines.append(f"**Type:** {data['meeting_type']}")
        lines.append(f"**App:** {data['meeting_app']}")
        lines.append("")

        # Summary
        if data.get("summary"):
            lines.append("## Summary")
            lines.append(data["summary"]["text"] or "")
            lines.append("")

            if data["summary"].get("key_points"):
                lines.append("### Key Points")
                for point in data["summary"]["key_points"]:
                    lines.append(f"- {point}")
                lines.append("")

            if data["summary"].get("decisions"):
                lines.append("### Decisions Made")
                for decision in data["summary"]["decisions"]:
                    lines.append(f"- {decision}")
                lines.append("")

        # Action Items
        if data.get("action_items"):
            lines.append("## Action Items")
            for item in data["action_items"]:
                status = "x" if item["status"] == "completed" else " "
                due = f" (Due: {item['due_date']})" if item.get("due_date") else ""
                lines.append(f"- [{status}] **{item['assignee']}**: {item['description']}{due}")
            lines.append("")

        # Conversations
        if data.get("conversations"):
            lines.append("## Transcript")
            for conv in data["conversations"]:
                speaker = "You" if conv["speaker"] == "user" else "Other"
                lines.append(f"**{speaker}:** {conv['heard_text']}")
                if conv.get("response_text"):
                    lines.append(f"> AI Suggestion: {conv['response_text']}")
                lines.append("")

        # Notes
        if data.get("notes"):
            lines.append("## Notes")
            lines.append(data["notes"])
            lines.append("")

        lines.append("---")
        lines.append("*Exported from ReadIn AI*")

        content = "\n".join(lines)
        filename = f"{data['title'].replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d')}.md"

        return (content.encode("utf-8"), filename, "text/markdown")

    def _export_json(self, data: Dict[str, Any]) -> tuple:
        """Export as JSON."""
        content = json.dumps(data, indent=2, default=str)
        filename = f"meeting_{data['id']}_{datetime.utcnow().strftime('%Y%m%d')}.json"
        return (content.encode("utf-8"), filename, "application/json")

    def _export_csv(self, data: Dict[str, Any]) -> tuple:
        """Export as CSV (conversations)."""
        output = io.StringIO()
        writer = csv.writer(output)

        # Header
        writer.writerow(["Timestamp", "Speaker", "Text", "AI Response"])

        # Conversations
        for conv in data.get("conversations", []):
            writer.writerow([
                conv["timestamp"],
                conv["speaker"],
                conv["heard_text"],
                conv.get("response_text", ""),
            ])

        content = output.getvalue()
        filename = f"meeting_{data['id']}_transcript.csv"
        return (content.encode("utf-8"), filename, "text/csv")

    def _export_html(self, data: Dict[str, Any]) -> tuple:
        """Export as HTML."""
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{data['title']}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #d4af37; }}
        h2 {{ color: #333; border-bottom: 1px solid #eee; padding-bottom: 10px; }}
        .meta {{ color: #666; margin-bottom: 20px; }}
        .action-item {{ padding: 10px; background: #f9f9f9; margin: 5px 0; border-radius: 4px; }}
        .action-item.completed {{ text-decoration: line-through; opacity: 0.7; }}
        .conversation {{ margin: 10px 0; }}
        .speaker {{ font-weight: bold; }}
        .ai-response {{ background: #fff3cd; padding: 10px; margin-top: 5px; border-radius: 4px; }}
        .summary {{ background: #e8f5e9; padding: 15px; border-radius: 4px; }}
        .footer {{ margin-top: 30px; padding-top: 20px; border-top: 1px solid #eee; color: #999; font-size: 12px; }}
    </style>
</head>
<body>
    <h1>{data['title']}</h1>
    <div class="meta">
        <p><strong>Date:</strong> {data['started_at']}</p>
        <p><strong>Duration:</strong> {data['duration_minutes']} minutes</p>
        <p><strong>Type:</strong> {data['meeting_type']} | <strong>App:</strong> {data['meeting_app']}</p>
    </div>
"""

        # Summary
        if data.get("summary"):
            html += f"""
    <h2>Summary</h2>
    <div class="summary">
        <p>{data['summary']['text'] or ''}</p>
    </div>
"""

        # Action Items
        if data.get("action_items"):
            html += "<h2>Action Items</h2>"
            for item in data["action_items"]:
                completed_class = "completed" if item["status"] == "completed" else ""
                due = f" (Due: {item['due_date']})" if item.get("due_date") else ""
                html += f'<div class="action-item {completed_class}"><strong>{item["assignee"]}:</strong> {item["description"]}{due}</div>'

        # Conversations
        if data.get("conversations"):
            html += "<h2>Transcript</h2>"
            for conv in data["conversations"]:
                speaker = "You" if conv["speaker"] == "user" else "Other"
                html += f'<div class="conversation"><span class="speaker">{speaker}:</span> {conv["heard_text"]}'
                if conv.get("response_text"):
                    html += f'<div class="ai-response">AI: {conv["response_text"]}</div>'
                html += "</div>"

        html += """
    <div class="footer">
        Exported from ReadIn AI
    </div>
</body>
</html>
"""

        filename = f"{data['title'].replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d')}.html"
        return (html.encode("utf-8"), filename, "text/html")

    def _export_pdf(self, data: Dict[str, Any]) -> tuple:
        """Export as PDF (requires reportlab or weasyprint)."""
        try:
            from weasyprint import HTML

            # Generate HTML first, then convert to PDF
            html_content, _, _ = self._export_html(data)

            pdf_content = HTML(string=html_content.decode("utf-8")).write_pdf()
            filename = f"{data['title'].replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d')}.pdf"

            return (pdf_content, filename, "application/pdf")
        except ImportError:
            # Fallback: return HTML with instructions
            html_content, _, _ = self._export_html(data)
            return (
                html_content,
                f"meeting_{data['id']}.html",
                "text/html",
            )

    def _export_docx(self, data: Dict[str, Any]) -> tuple:
        """Export as Word document (requires python-docx)."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(data["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Meta info
            doc.add_paragraph(f"Date: {data['started_at']}")
            doc.add_paragraph(f"Duration: {data['duration_minutes']} minutes")
            doc.add_paragraph(f"Type: {data['meeting_type']} | App: {data['meeting_app']}")

            # Summary
            if data.get("summary"):
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(data["summary"]["text"] or "")

            # Action Items
            if data.get("action_items"):
                doc.add_heading("Action Items", level=1)
                for item in data["action_items"]:
                    status = "[x]" if item["status"] == "completed" else "[ ]"
                    due = f" (Due: {item['due_date']})" if item.get("due_date") else ""
                    doc.add_paragraph(f"{status} {item['assignee']}: {item['description']}{due}")

            # Conversations
            if data.get("conversations"):
                doc.add_heading("Transcript", level=1)
                for conv in data["conversations"]:
                    speaker = "You" if conv["speaker"] == "user" else "Other"
                    doc.add_paragraph(f"{speaker}: {conv['heard_text']}")

            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            filename = f"{data['title'].replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
            return (
                output.getvalue(),
                filename,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except ImportError:
            # Fallback to Markdown
            return self._export_markdown(data)

    def _export_meetings_csv(self, meetings_data: List[Dict]) -> tuple:
        """Export multiple meetings as CSV summary."""
        output = io.StringIO()
        writer = csv.writer(output)

        # Header
        writer.writerow([
            "ID", "Title", "Type", "App", "Date", "Duration (min)",
            "Status", "Conversations", "Action Items"
        ])

        # Rows
        for data in meetings_data:
            writer.writerow([
                data["id"],
                data["title"],
                data["meeting_type"],
                data["meeting_app"],
                data["started_at"],
                data["duration_minutes"],
                data["status"],
                len(data.get("conversations", [])),
                len(data.get("action_items", [])),
            ])

        content = output.getvalue()
        filename = f"meetings_export_{datetime.utcnow().strftime('%Y%m%d')}.csv"
        return (content.encode("utf-8"), filename, "text/csv")

    def _export_meetings_markdown(self, meetings_data: List[Dict]) -> tuple:
        """Export multiple meetings as Markdown."""
        lines = ["# Meeting Export", "", f"*Exported on {datetime.utcnow().strftime('%Y-%m-%d')}*", ""]

        for data in meetings_data:
            lines.append(f"## {data['title']}")
            lines.append(f"**Date:** {data['started_at']} | **Duration:** {data['duration_minutes']} min")
            lines.append("")

            if data.get("summary"):
                lines.append(data["summary"]["text"] or "No summary available")
                lines.append("")

            lines.append("---")
            lines.append("")

        content = "\n".join(lines)
        filename = f"meetings_export_{datetime.utcnow().strftime('%Y%m%d')}.md"
        return (content.encode("utf-8"), filename, "text/markdown")
